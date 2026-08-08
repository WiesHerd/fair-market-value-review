#!/usr/bin/env python3
"""
CV-Only Market Anchor Builder
For when you only have a candidate's CV (no offer amount, no internal equity file).
Produces a market-anchored recommended range + exhibit.

USAGE:
  python3 build_cv_only_market_anchor.py \
    --name "Lastname, Firstname" \
    --specialty "Pediatric Critical Care" \
    --yoe 8 \
    --survey-file survey_combined.xlsx \
    --output-dir ./output

Output:
  - market_anchor_exhibit.png  (horizontal percentile band chart)
  - market_anchor_report.docx   (2-page market justification)
  - market_anchor_report.pdf    (PDF export)
  - market_anchor_summary.json  (machine-readable summary)
"""
import argparse
import json
import sys
import subprocess
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import sys as _sys
    from pathlib import Path as _Path
    _sys.path.insert(0, str(_Path(__file__).parent))
    from docx_style_helpers import (
        style_header_row, band_and_border_rows, right_align_columns,
        remove_table_grid_borders, add_page_number_footer, set_page_margins,
        format_title_block, style_section_heading, add_kpi_cards, add_callout,
    )
except ImportError:
    print("ERROR: python-docx required. Install: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
except ImportError:
    print("ERROR: matplotlib required. Install: pip install matplotlib", file=sys.stderr)
    sys.exit(1)

# Import shared functions from fmv_workbook_generator
sys.path.insert(0, str(Path(__file__).parent))
from fmv_workbook_generator import (
    resolve_dec,
    lookup_benchmarks_from_survey,
    parse_benchmarks,
    SURVEY_SHEET_NAME,
)


# ─── Percentile interpolation ────────────────────────────────────────

def interpolate_percentile(value, p25, p50, p75, p90):
    """Piecewise linear interpolation to estimate percentile position.
    Returns float percentile or None if inputs missing.
    """
    if not all(v is not None and v > 0 for v in [p25, p50, p75, p90]):
        return None
    if value <= p25:
        return 25.0 * (value / p25) if p25 > 0 else 0
    elif value <= p50:
        return 25.0 + 25.0 * (value - p25) / (p50 - p25)
    elif value <= p75:
        return 50.0 + 25.0 * (value - p50) / (p75 - p50)
    elif value <= p90:
        return 75.0 + 15.0 * (value - p75) / (p90 - p75)
    else:
        return 90.0 + 10.0 * (value - p90) / p90  # extrapolation above p90


def blended_benchmark(benchmarks, section='tcc'):
    """Compute blended weighted-average percentiles across all sources.
    Returns dict with p25, p50, p75, p90 (blended).
    """
    section_data = benchmarks.get(section, {})
    blended = {}
    for pct in ['25', '50', '75', '90']:
        weighted_sum = 0
        total_n = 0
        for source_name, data in section_data.items():
            n = data.get('n', 0)
            val = data.get(pct)
            if n > 0 and val and val > 0:
                weighted_sum += val * n
                total_n += n
        blended[pct] = (weighted_sum / total_n) if total_n > 0 else None
    return blended


# ─── Exhibit ──────────────────────────────────────────────────────────

def build_exhibit(specialty, blended_tcc, recommended_low, recommended_high,
                  recommended_mid, output_path):
    """Build a horizontal percentile band chart showing the recommended range
    positioned against the market distribution.
    """
    fig, ax = plt.subplots(figsize=(10, 4))

    p25 = blended_tcc.get('25', 0)
    p50 = blended_tcc.get('50', 0)
    p75 = blended_tcc.get('75', 0)
    p90 = blended_tcc.get('90', 0)

    # Market band
    ax.barh(0, p75 - p25, left=p25, height=0.4, color='#2F7DB8', alpha=0.25, label='p25–p75 band')
    ax.plot([p50, p50], [-0.25, 0.25], color='#00485F', linewidth=3, label=f'p50 (${p50:,.0f})')
    ax.plot(p90, 0, 'v', color='#F0832E', markersize=10, label=f'p90 (${p90:,.0f})')

    # Recommended range
    ax.barh(-0.5, recommended_high - recommended_low, left=recommended_low,
            height=0.3, color='#F0832E', alpha=0.3, label=f'Recommended range')
    ax.plot(recommended_mid, -0.5, 'D', color='#F0832E', markersize=10,
            label=f'Recommended midpoint (${recommended_mid:,.0f})')

    ax.set_ylim(-1, 0.5)
    ax.set_yticks([0, -0.5])
    ax.set_yticklabels(['Market Distribution', 'Recommended Range'], fontsize=11)
    ax.set_xlabel('Total Cash Compensation ($)', fontsize=11)
    ax.set_title(f'Market Anchor Analysis — {specialty}', fontsize=13, fontweight='bold')
    ax.legend(loc='lower right', fontsize=9)
    ax.grid(axis='x', alpha=0.3)
    ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f'${x/1000:.0f}K'))

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    return output_path


# ─── Report ──────────────────────────────────────────────────────────

def build_report(name, specialty, yoe, blended_tcc, recommended_range, exhibit_path,
                 output_docx, survey_sources=None):
    """Build a 2-page market-anchored DOCX report for CV-only intake.
    """
    doc = Document()
    set_page_margins(doc, 0.6)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Title ──
    h = doc.add_heading('Market Anchor Analysis', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Provider: {name}").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Specialty: {specialty} | YOE: {yoe}")

    doc.add_paragraph()

    # ── Caveat ──
    add_callout(
        doc,
        'This analysis is based solely on market survey data and the candidate\'s CV. '
        'No internal equity cohort file or existing offer amount was available. '
        'Recommendations should be validated against internal equity before finalizing.',
        kind='warning', lead='IMPORTANT CAVEAT —')
    doc.add_paragraph()

    # ── Market Benchmarks ──
    doc.add_heading('Blended Market Benchmarks (TCC)', level=1)

    table = doc.add_table(rows=1, cols=3, style='Table Grid')
    hdr = table.rows[0].cells
    for i, h_text in enumerate(['Percentile', 'Blended TCC', 'Source']):
        hdr[i].text = h_text
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for pct, label in [('25', '25th'), ('50', '50th (Median)'), ('75', '75th'), ('90', '90th')]:
        row = table.add_row().cells
        row[0].text = label
        val = blended_tcc.get(pct)
        row[1].text = f"${val:,.0f}" if val else "N/A"
        row[2].text = 'Blended weighted-average' if survey_sources else 'Single source'

    remove_table_grid_borders(table)
    style_header_row(table)
    band_and_border_rows(table)
    right_align_columns(table, [1])

    # ── Recommended Range ──
    doc.add_heading('Recommended Range', level=1)

    low, mid, high = recommended_range
    add_kpi_cards(doc, [
        ("Low (≈p25–p30)", f"${low:,.0f}", "conservative entry"),
        ("Midpoint (≈p35)", f"${mid:,.0f}", "directional estimate"),
        ("High (≈p40–p50)", f"${high:,.0f}", "upper governance bound"),
        ("Market p50", f"${blended_tcc.get('50') or 0:,.0f}", "blended median"),
    ])
    doc.add_paragraph()
    doc.add_paragraph(f"Low end (≈p25–p30): ${low:,.0f}")
    doc.add_paragraph(f"Midpoint (≈p35, directional): ${mid:,.0f}")
    doc.add_paragraph(f"High end (≈p40–p50): ${high:,.0f}")

    p = doc.add_paragraph()
    p.add_run('Methodology: ').bold = True
    p.add_run(
        'Recommended range targets the p25–p50 band for a candidate at this YOE level, '
        'using piecewise linear interpolation between survey percentile points. '
        'Midpoint is a directional estimate at approximately p35, rounded to the nearest $1,000.'
    )

    # ── Exhibit ──
    doc.add_heading('Market Position Exhibit', level=1)
    if exhibit_path and Path(exhibit_path).exists():
        doc.add_picture(exhibit_path, width=Inches(7.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Footnotes ──
    doc.add_heading('Footnotes', level=1)
    doc.add_paragraph('Survey distributions are not linear between percentile points; interpolated estimates are directional, not exact.', style='List Bullet')
    doc.add_paragraph('Blended benchmark = weighted average across aggregated survey data sources by n-count.', style='List Bullet')
    doc.add_paragraph('No internal equity or cost impact analysis is included in this CV-only assessment.', style='List Bullet')
    doc.add_paragraph(f'Analysis date: {datetime.now().strftime("%Y-%m-%d")}', style='List Bullet')

    add_page_number_footer(doc, "Market Anchor Analysis — Confidential")

    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            style_section_heading(para)

    doc.save(output_docx)
    return output_docx


# ─── PDF Export ────────────────────────────────────────────────────────

def convert_to_pdf(docx_path, output_dir=None):
    """Convert DOCX to PDF using LibreOffice headless."""
    output_dir = output_dir or str(Path(docx_path).parent)
    try:
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf',
             '--outdir', output_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        pdf_path = docx_path.replace('.docx', '.pdf')
        if Path(pdf_path).exists():
            return pdf_path
        return None
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return None


# ─── Main ─────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Build CV-Only Market Anchor Report')
    parser.add_argument('--name', required=True, help='Provider name')
    parser.add_argument('--specialty', required=True, help='Specialty (must match survey file)')
    parser.add_argument('--yoe', type=int, default=0, help='Years of experience')
    parser.add_argument('--survey-file', help='Survey Combined Excel file')
    parser.add_argument('--benchmarks', help='Benchmark text file (fallback)')
    parser.add_argument('--output-dir', default='.', help='Output directory')
    parser.add_argument('--no-pdf', action='store_true')

    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Get benchmarks
    if args.survey_file:
        print(f"Looking up benchmarks for '{args.specialty}' in {args.survey_file}...")
        benchmarks = lookup_benchmarks_from_survey(args.survey_file, args.specialty)
    elif args.benchmarks:
        with open(args.benchmarks) as f:
            benchmarks = parse_benchmarks(f.read())
    else:
        print("ERROR: --survey-file or --benchmarks required", file=sys.stderr)
        sys.exit(1)

    # Compute blended TCC
    blended_tcc = blended_benchmark(benchmarks, 'tcc')
    p25 = blended_tcc.get('25')
    p50 = blended_tcc.get('50')
    p75 = blended_tcc.get('75')
    p90 = blended_tcc.get('90')

    if not p25 or not p50:
        print("ERROR: Insufficient benchmark data for TCC percentiles", file=sys.stderr)
        sys.exit(1)

    print(f"Blended TCC: p25=${p25:,.0f}, p50=${p50:,.0f}, p75=${p75:,.0f}, p90=${p90 or 0:,.0f}")

    # Recommended range: target p25-p50, midpoint ~p35
    recommended_low = round(p25 / 1000) * 1000
    recommended_mid = round((p25 + 0.35 * (p50 - p25)) / 1000) * 1000
    recommended_high = round((p25 + 0.6 * (p50 - p25)) / 1000) * 1000

    print(f"Recommended range: ${recommended_low:,.0f} – ${recommended_high:,.0f} (mid: ${recommended_mid:,.0f})")

    # Build exhibit
    exhibit_path = str(output_dir / 'market_anchor_exhibit.png')
    build_exhibit(args.specialty, blended_tcc, recommended_low, recommended_high,
                  recommended_mid, exhibit_path)
    print(f"Exhibit: {exhibit_path}")

    # Build report
    docx_path = str(output_dir / 'market_anchor_report.docx')
    sources_with_data = [name for name, data in benchmarks.get('tcc', {}).items()
                         if data.get('n', 0) > 0]
    build_report(args.name, args.specialty, args.yoe, blended_tcc,
                 (recommended_low, recommended_mid, recommended_high),
                 exhibit_path, docx_path, survey_sources=sources_with_data)
    print(f"DOCX: {docx_path}")

    # PDF
    if not args.no_pdf:
        pdf_path = convert_to_pdf(docx_path)
        if pdf_path:
            print(f"PDF: {pdf_path}")

    # Summary JSON
    summary = {
        'name': args.name,
        'specialty': args.specialty,
        'yoe': args.yoe,
        'blended_tcc': blended_tcc,
        'recommended_range': {
            'low': recommended_low,
            'mid': recommended_mid,
            'high': recommended_high,
        },
        'analysis_date': datetime.now().isoformat(),
        'sources': sources_with_data,
        'caveat': 'CV-only assessment — no internal equity or cost impact analysis.',
    }
    summary_path = str(output_dir / 'market_anchor_summary.json')
    with open(summary_path, 'w') as f:
        json.dump(summary, f, indent=2)
    print(f"Summary: {summary_path}")


if __name__ == '__main__':
    main()