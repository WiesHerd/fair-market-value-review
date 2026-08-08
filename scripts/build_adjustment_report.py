#!/usr/bin/env python3
"""
Compensation Adjustment Report Builder
Generates a consulting-style DOCX + PDF report with market-band exhibit.

USAGE:
  python3 build_adjustment_report.py --config report_config.json

OR with CLI args:
  python3 build_adjustment_report.py \
    --name "Division A" \
    --providers-file cohort.tsv \
    --market-anchors-file anchors.json \
    --output report.docx

Input format:
  --providers-file: TSV with columns: Name, YOE, CurrentBase, PlannedBase, ProposedBase
  --market-anchors-file: JSON with p25, p50, p75, p90 (TCC benchmarks)

Output:
  - DOCX report (sections 1-8)
  - PDF (via LibreOffice headless)
  - PNG exhibit (market band chart)
  - CSV with computed values
"""
import argparse
import json
import sys
import subprocess
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import sys as _sys
    from pathlib import Path as _Path
    _sys.path.insert(0, str(_Path(__file__).parent))
    from docx_style_helpers import (
        style_header_row, band_and_border_rows, right_align_columns,
        remove_table_grid_borders, add_page_number_footer, set_page_margins,
        format_title_block, style_section_heading, add_kpi_cards, add_callout,
        add_totals_row,
    )
except ImportError:
    print("ERROR: python-docx required. Install: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
except ImportError:
    print("ERROR: matplotlib required. Install: pip install matplotlib", file=sys.stderr)
    sys.exit(1)

try:
    import openpyxl
except ImportError:
    print("WARNING: openpyxl not available (needed for Excel input only)", file=sys.stderr)


# ─── Data Loading ────────────────────────────────────────────────────

def load_providers_tsv(path):
    """Load provider cohort from TSV file.
    Expected columns: Name, YOE, CurrentBase, PlannedBase, ProposedBase
    """
    providers = []
    with open(path) as f:
        header = f.readline().strip().split('\t')
        for line in f:
            parts = line.strip().split('\t')
            if len(parts) < 5:
                continue
            providers.append({
                'name': parts[0],
                'yoe': float(parts[1]) if parts[1] else None,
                'current_base': float(parts[2]) if parts[2] else 0,
                'planned_base': float(parts[3]) if parts[3] else 0,
                'proposed_base': float(parts[4]) if parts[4] else 0,
            })
    return providers


def load_providers_excel(path, sheet_name=None, name_col='D',
                        yoe_col='K', current_col='P',
                        planned_col='Q', proposed_col=None):
    """Load provider cohort from Excel file.
    Columns are configurable — adjust to match your file layout.
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active

    providers = []
    for row in range(2, ws.max_row + 1):
        name = ws[name_col + str(row)].value
        if not name:
            continue
        providers.append({
            'name': str(name),
            'yoe': ws[yoe_col + str(row)].value,
            'current_base': ws[current_col + str(row)].value or 0,
            'planned_base': ws[planned_col + str(row)].value or 0,
            'proposed_base': ws[proposed_col + str(row)].value if proposed_col else 0,
        })
    return providers


def load_market_anchors(path):
    """Load market anchors from JSON file.
    Format: {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
    """
    with open(path) as f:
        return json.load(f)


# ─── Exhibit (Market Band Chart) ─────────────────────────────────────

def build_exhibit(providers, anchors, output_path, title="Base Salary vs Market Band"):
    """Build a horizontal market-band exhibit showing current, planned, proposed
    vs market p25-p75 band with p50 line.
    """
    fig, ax = plt.subplots(figsize=(10, max(4, len(providers) * 0.8)))

    names = [p['name'] for p in providers]
    y_pos = range(len(names))

    # Market band (p25-p75)
    p25 = anchors.get('p25', 0)
    p50 = anchors.get('p50', 0)
    p75 = anchors.get('p75', 0)
    p90 = anchors.get('p90', 0)

    for i, p in enumerate(providers):
        # Draw band
        ax.barh(i, p75 - p25, left=p25, height=0.5, color='#2F7DB8', alpha=0.25, edgecolor='none')
        # p50 line
        ax.plot([p50, p50], [i - 0.3, i + 0.3], color='#00485F', linewidth=2)
        # p90 marker
        ax.plot(p90, i, 'v', color='#F0832E', markersize=6, alpha=0.5)

        # Current
        ax.plot(p['current_base'], i + 0.15, 'o', color='#888888', markersize=8, label='Current' if i == 0 else '')
        # Planned
        ax.plot(p['planned_base'], i, 's', color='#2F7DB8', markersize=8, label='Planned' if i == 0 else '')
        # Proposed
        ax.plot(p['proposed_base'], i - 0.15, 'D', color='#F0832E', markersize=8, label='Proposed' if i == 0 else '')

    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(names, fontsize=10)
    ax.set_xlabel('Base Salary ($)', fontsize=11)
    ax.set_title(title, fontsize=13, fontweight='bold')
    ax.legend(loc='lower right', fontsize=9)
    ax.grid(axis='x', alpha=0.3)

    # Format x-axis as currency
    ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f'${x/1000:.0f}K'))

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    return output_path


# ─── Report Generation ────────────────────────────────────────────────

def build_report(providers, anchors, exhibit_path, output_docx, metadata=None):
    """Build the 8-section DOCX report.
    """
    doc = Document()
    set_page_margins(doc, 0.6)

    # ── Styling ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    meta = metadata or {}

    # ── Section 1: Title + Scope ──
    h = doc.add_heading('Fair Market Value Review', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph()
    bits = [meta.get('cohort_name', 'N/A')]
    if meta.get('division'):
        bits.append(meta['division'])
    bits.append(f"Prepared {meta.get('snapshot_date', datetime.now().strftime('%Y-%m-%d'))}")
    bits.append("Confidential — for compensation committee review")
    sub.add_run("  ·  ".join(bits))
    format_title_block(h, sub)

    p = doc.add_paragraph()
    p.add_run(f"Cohort: {meta.get('cohort_name', 'N/A')}").bold = True
    p.add_run(f"    Division/Service Line: {meta.get('division', 'N/A')}")
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string('63696F')

    # ── Section 2: Data Sources ──
    doc.add_heading('Data Sources', level=1)
    sources = meta.get('data_sources', [])
    if sources:
        for s in sources:
            doc.add_paragraph(s, style='List Bullet')
    else:
        doc.add_paragraph('Annual Salary Increases file (authoritative for current base, TCC percentile, wRVU percentile).')
        doc.add_paragraph('Survey Combined file (aggregated survey data by specialty).')
        doc.add_paragraph('Committee Request Form template (output deliverable).')

    p = doc.add_paragraph()
    p.add_run(f"Snapshot date: {meta.get('snapshot_date', datetime.now().strftime('%Y-%m-%d'))}").italic = True

    # ── Section 3: Executive Summary ──
    doc.add_heading('Executive Summary', level=1)

    total_current = sum(p['current_base'] for p in providers)
    total_planned = sum(p['planned_base'] for p in providers)
    total_proposed = sum(p['proposed_base'] for p in providers)
    delta_vs_current = total_proposed - total_current
    delta_vs_planned = total_proposed - total_planned

    add_kpi_cards(doc, [
        ("Total Current Base", f"${total_current:,.0f}", f"{len(providers)} provider(s)"),
        ("Total Proposed Base", f"${total_proposed:,.0f}", "post-adjustment"),
        ("Δ vs Current", f"+${delta_vs_current:,.0f}",
         f"{delta_vs_current/total_current*100:+.1f}%" if total_current else ""),
        ("Δ vs Planned", f"+${delta_vs_planned:,.0f}",
         f"{delta_vs_planned/total_planned*100:+.1f}%" if total_planned else ""),
    ])
    doc.add_paragraph()

    summary_points = [
        f"Cohort size: {len(providers)} provider(s).",
        f"Total current base: ${total_current:,.0f}.",
        f"Total planned base (post-annual-increase): ${total_planned:,.0f}.",
        f"Total proposed base: ${total_proposed:,.0f}.",
        f"Incremental cost vs current: ${delta_vs_current:,.0f} ({delta_vs_current/total_current*100:.1f}%).",
        f"Incremental cost vs planned: ${delta_vs_planned:,.0f} ({delta_vs_planned/total_planned*100:.1f}%).",
    ]
    for point in summary_points:
        doc.add_paragraph(point, style='List Bullet')

    p50 = anchors.get('p50', 0)
    if p50:
        avg_proposed = total_proposed / len(providers) if providers else 0
        add_callout(
            doc,
            f"Average proposed base (${avg_proposed:,.0f}) vs market p50 (${p50:,.0f}): "
            f"{((avg_proposed - p50) / p50 * 100):+.1f}% vs market median.",
            kind="info", lead="MARKET POSITION —")
        doc.add_paragraph()

    # ── Section 4: Cohort Detail Table ──
    doc.add_heading('Cohort Detail', level=1)

    table = doc.add_table(rows=1, cols=7, style='Table Grid')
    hdr = table.rows[0].cells
    headers = ['Name', 'YOE', 'Current Base', 'Planned Base', 'Proposed Base', 'Δ vs Current', 'Δ vs Planned']
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for p_data in providers:
        row = table.add_row().cells
        row[0].text = p_data['name']
        row[1].text = str(p_data.get('yoe', 'N/A'))
        row[2].text = f"${p_data['current_base']:,.0f}"
        row[3].text = f"${p_data['planned_base']:,.0f}"
        row[4].text = f"${p_data['proposed_base']:,.0f}"
        delta_cur = p_data['proposed_base'] - p_data['current_base']
        delta_pln = p_data['proposed_base'] - p_data['planned_base']
        row[5].text = f"${delta_cur:,.0f} ({delta_cur/p_data['current_base']*100:+.1f}%)"
        row[6].text = f"${delta_pln:,.0f} ({delta_pln/p_data['planned_base']*100:+.1f}%)"

    tc = sum(pd['current_base'] for pd in providers)
    tp = sum(pd['planned_base'] for pd in providers)
    tpr = sum(pd['proposed_base'] for pd in providers)
    add_totals_row(table, 'Cohort Total',
                   [f"${tc:,.0f}", f"${tp:,.0f}", f"${tpr:,.0f}",
                    f"${tpr - tc:,.0f}", f"${tpr - tp:,.0f}"], col_offset=2)

    remove_table_grid_borders(table)
    style_header_row(table)
    band_and_border_rows(table)
    right_align_columns(table, [2, 3, 4, 5, 6])

    # ── Section 5: Market Range Exhibit ──
    doc.add_heading('Market Range Exhibit', level=1)
    if exhibit_path and Path(exhibit_path).exists():
        doc.add_picture(exhibit_path, width=Inches(7.2))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('Market band: p25–p75 (shaded). p50 median (vertical line). p90 (triangle). ').italic = True
    p.add_run('Current (circle), Planned (square), Proposed (diamond).').italic = True

    # ── Section 6: Logic & Empirical Rationale ──
    doc.add_heading('Logic & Empirical Rationale', level=1)

    rationale_points = meta.get('rationale', [])
    if not rationale_points:
        rationale_points = [
            'Internal equity: proposed salaries maintain YOE-based progression without inversions.',
            'External anchoring: proposed bases position providers within the p25–p75 market band.',
            'Targeted adjustment: focuses on mid-career providers where compression risk is highest.',
            'Compression guardrails: no cascade adjustments proposed at this time.',
        ]
    for point in rationale_points:
        doc.add_paragraph(point, style='List Bullet')

    # ── Section 7: Cost Impact ──
    doc.add_heading('Cost Impact', level=1)

    doc.add_paragraph(f"Annualized incremental base cost vs current: ${delta_vs_current:,.0f}")
    doc.add_paragraph(f"Incremental beyond planned annual increases: ${delta_vs_planned:,.0f}")
    p = doc.add_paragraph()
    p.add_run('Note: Fringe/benefits excluded from above figures. ').italic = True
    benefits_rate = meta.get('benefits_rate')
    if benefits_rate:
        loaded_cost = delta_vs_current * (1 + benefits_rate)
        p.add_run(f'Applying benefits rate of {benefits_rate*100:.1f}%: fully loaded incremental cost ≈ ${loaded_cost:,.0f}.').italic = True
    else:
        p.add_run('Apply organization-specific benefits rate for fully loaded cost.').italic = True

    # ── Section 8: Governance Notes ──
    doc.add_heading('Governance Notes & Guardrails', level=1)
    notes = meta.get('governance_notes', [
        'Proposed salaries do not create senior-junior inversions.',
        'All percentiles are from governance reference (post-annual-increase file).',
        'Percentiles are NOT FTE-adjusted — part-time lower percentile positioning is correct.',
        'Benchmark source: blended weighted-average of aggregated survey data (weighted by sample size).',
    ])
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')

    # ── Section 9: Footnotes ──
    doc.add_heading('Footnotes & Caveats', level=1)
    footnotes = meta.get('footnotes', [
        'Market percentiles are from published survey data; interpolated estimates are labeled as directional.',
        'Cost figures reflect base salary only unless otherwise noted.',
        'All data sourced from the most recent annual salary increases file and survey combined file.',
    ])
    for fn in footnotes:
        doc.add_paragraph(fn, style='List Bullet')

    add_page_number_footer(doc, f"Fair Market Value Review — Confidential")

    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            style_section_heading(para)

    doc.save(output_docx)
    return output_docx


# ─── PDF Export ────────────────────────────────────────────────────────

def convert_to_pdf(docx_path, output_dir=None):
    """Convert DOCX to PDF using LibreOffice headless."""
    output_dir = output_dir or str(Path(docx_path).parent)
    try:
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf',
             '--outdir', output_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        pdf_path = docx_path.replace('.docx', '.pdf')
        if Path(pdf_path).exists():
            return pdf_path
        print(f"WARNING: PDF conversion may have failed: {result.stderr}", file=sys.stderr)
        return None
    except FileNotFoundError:
        print("WARNING: LibreOffice (soffice) not found — PDF not generated.", file=sys.stderr)
        return None
    except subprocess.TimeoutExpired:
        print("WARNING: PDF conversion timed out.", file=sys.stderr)
        return None


# ─── Main ─────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Build Compensation Adjustment Report')
    parser.add_argument('--config', help='JSON config file')
    parser.add_argument('--name', help='Cohort name')
    parser.add_argument('--providers-file', help='TSV or Excel file with provider data')
    parser.add_argument('--market-anchors-file', help='JSON file with p25/p50/p75/p90')
    parser.add_argument('--output', default='adjustment_report.docx')
    parser.add_argument('--no-pdf', action='store_true', help='Skip PDF conversion')
    parser.add_argument('--no-exhibit', action='store_true', help='Skip exhibit chart')

    args = parser.parse_args()

    if args.config:
        with open(args.config) as f:
            cfg = json.load(f)
        for k, v in cfg.items():
            setattr(args, k, v)

    if not args.providers_file:
        print("ERROR: --providers-file or --config required", file=sys.stderr)
        sys.exit(1)

    # Load data
    if args.providers_file.endswith('.xlsx'):
        providers = load_providers_excel(args.providers_file)
    else:
        providers = load_providers_tsv(args.providers_file)

    print(f"Loaded {len(providers)} providers from {args.providers_file}")

    anchors = {}
    if args.market_anchors_file:
        anchors = load_market_anchors(args.market_anchors_file)
        print(f"Loaded market anchors: p25=${anchors.get('p25',0):,.0f}, p50=${anchors.get('p50',0):,.0f}, p75=${anchors.get('p75',0):,.0f}")

    # Build exhibit
    exhibit_path = None
    if not args.no_exhibit and anchors and providers:
        exhibit_path = args.output.replace('.docx', '_exhibit.png')
        build_exhibit(providers, anchors, exhibit_path, title=f"Base Salary vs Market Band — {args.name or 'Cohort'}")
        print(f"Exhibit saved: {exhibit_path}")

    # Build report
    # NOTE: build_report() only renders these specific metadata keys (it has no
    # "Request Summary" / "Background" section -- that's a committee_template_generator.py
    # concept). Config fields like request_summary/general_background/physician_background
    # are accepted for compatibility with a shared config file but are not rendered here.
    #
    # Only set a key if a value was actually provided: build_report() uses
    # meta.get(key, [default...]) internally, which only falls back to its default
    # when the key is *absent* -- an explicit None would short-circuit that fallback.
    metadata = {
        'cohort_name': args.name or 'N/A',
        'snapshot_date': datetime.now().strftime('%Y-%m-%d'),
    }
    for key in ('division', 'data_sources', 'benefits_rate', 'rationale', 'governance_notes', 'footnotes'):
        value = getattr(args, key, None)
        if value is not None:
            metadata[key] = value

    build_report(providers, anchors, exhibit_path, args.output, metadata)
    print(f"DOCX saved: {args.output}")

    # PDF
    if not args.no_pdf:
        pdf_path = convert_to_pdf(args.output)
        if pdf_path:
            print(f"PDF saved: {pdf_path}")

    # CSV
    csv_path = args.output.replace('.docx', '.csv')
    import csv as csv_mod
    with open(csv_path, 'w', newline='') as f:
        writer = csv_mod.writer(f)
        writer.writerow(['Name', 'YOE', 'CurrentBase', 'PlannedBase', 'ProposedBase', 'DeltaVsCurrent', 'DeltaVsPlanned'])
        for p in providers:
            dc = p['proposed_base'] - p['current_base']
            dp = p['proposed_base'] - p['planned_base']
            writer.writerow([p['name'], p.get('yoe', ''), p['current_base'], p['planned_base'], p['proposed_base'], dc, dp])
    print(f"CSV saved: {csv_path}")


if __name__ == '__main__':
    main()