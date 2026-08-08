"""
Shared DOCX styling helpers -- python-docx doesn't expose cell shading or
footer page-number fields directly, so these wrap the underlying OOXML.
Used by build_adjustment_report.py and build_cv_only_market_anchor.py so
both reports get the same "designed" look: shaded header row, banded rows,
right-aligned numeric columns, and a real page-number footer.
"""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = "00485F"
BLUE = "2F7DB8"
ORANGE = "F0832E"
BAND = "F5F8FA"
BORDER_GREY = "D9DEE3"


def shade_cell(cell, hex_color):
    """Set a table cell's background fill color (no python-docx API for this)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=BORDER_GREY, sz=4):
    """Thin bottom border only (modern report look, not a full grid)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('bottom',):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def style_header_row(table, bg_color=NAVY, font_color="FFFFFF"):
    """Bold white text on a dark filled header row -- the single highest-impact
    change for making a python-docx table look designed instead of default."""
    header = table.rows[0]
    for cell in header.cells:
        shade_cell(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(font_color)
                run.font.size = Pt(10)


def band_and_border_rows(table, band_color=BAND):
    """Light banding on alternating data rows + a thin bottom border on every
    row, instead of the default heavy all-cell black grid."""
    for i, row in enumerate(table.rows[1:]):
        for cell in row.cells:
            if i % 2 == 1:
                shade_cell(cell, band_color)
            set_cell_borders(cell)


def right_align_columns(table, col_indices, skip_header=True):
    start = 1 if skip_header else 0
    for row in table.rows[start:]:
        for idx in col_indices:
            if idx < len(row.cells):
                for paragraph in row.cells[idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def remove_table_grid_borders(table):
    """The default 'Table Grid' style draws a border on every cell edge. Strip
    it so our custom bottom-border-only styling (set_cell_borders) is what
    actually shows."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        borders.append(el)
    tblPr.append(borders)


def add_page_number_footer(doc, label):
    """Left: label text. Right: 'Page X of Y' using real Word field codes
    (so it updates correctly if the document is edited/repaginated)."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    tab_stops = p.paragraph_format.tab_stops
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches
    tab_stops.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)

    run = p.add_run(label + "\t")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("9AA5AD")

    run2 = p.add_run("Page ")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor.from_string("9AA5AD")

    def field(instr):
        r = OxmlElement('w:r')
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr_el = OxmlElement('w:instrText')
        instr_el.set(qn('xml:space'), 'preserve')
        instr_el.text = instr
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        r.append(fld_begin)
        r.append(instr_el)
        r.append(fld_sep)
        r.append(fld_end)
        return r

    p._p.append(field('PAGE'))
    run3 = p.add_run(" of ")
    run3.font.size = Pt(8)
    run3.font.color.rgb = RGBColor.from_string("9AA5AD")
    p._p.append(field('NUMPAGES'))


def set_page_margins(doc, inches=0.6):
    """Set uniform page margins (default 0.6in on all sides)."""
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


# ─── FMV brand kit (matches the Excel ledger design) ─────────────────────
FMV_NAVY = "1F3A5F"
FMV_BURGUNDY = "8B3A3A"
FMV_SLATE = "63696F"
FMV_LIGHT_NAVY = "EEF2F7"
FMV_LIGHT_BURGUNDY = "F7EEEE"


def _para_bottom_border(paragraph, color=FMV_NAVY, sz=12):
    """Add a bottom rule to a paragraph (python-docx has no API for this)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def format_title_block(title_heading, subtitle_paragraph=None):
    """Serif navy display title + navy rule under the subtitle -- mirrors the
    Excel template's identity so the packet reads as one branded set."""
    for run in title_heading.runs:
        run.font.name = 'Georgia'
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(FMV_NAVY)
    if subtitle_paragraph is not None:
        for run in subtitle_paragraph.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor.from_string(FMV_SLATE)
        _para_bottom_border(subtitle_paragraph, color=FMV_NAVY, sz=16)


def style_section_heading(heading):
    """Small-caps navy caption with a rule. Uses small_caps rather than
    uppercasing the text, so heading strings stay machine-readable."""
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.small_caps = True
        run.font.color.rgb = RGBColor.from_string(FMV_NAVY)
    _para_bottom_border(heading, color=FMV_NAVY, sz=12)


def add_kpi_cards(doc, cards):
    """Row of KPI metric cards: [(label, value, sub), ...]."""
    table = doc.add_table(rows=1, cols=len(cards))
    table.autofit = True
    remove_table_grid_borders(table)
    fills = [FMV_LIGHT_NAVY, FMV_LIGHT_BURGUNDY]
    values = [FMV_NAVY, FMV_BURGUNDY]
    for i, (label, value, sub) in enumerate(cards):
        cell = table.rows[0].cells[i]
        shade_cell(cell, fills[i % 2])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label.upper() + "\n")
        r.font.size = Pt(7.5); r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(FMV_SLATE)
        r2 = p.add_run(value)
        r2.font.size = Pt(16); r2.font.bold = True
        r2.font.color.rgb = RGBColor.from_string(values[i % 2])
        if sub:
            r3 = p.add_run("\n" + sub)
            r3.font.size = Pt(7.5)
            r3.font.color.rgb = RGBColor.from_string(FMV_SLATE)
    return table


def add_callout(doc, text, kind="info", lead=None):
    """Single-cell callout with a thick left accent bar and light fill."""
    color = FMV_NAVY if kind == "info" else FMV_BURGUNDY
    fill = FMV_LIGHT_NAVY if kind == "info" else FMV_LIGHT_BURGUNDY
    table = doc.add_table(rows=1, cols=1)
    remove_table_grid_borders(table)
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24'); left.set(qn('w:color'), color)
    borders.append(left); tcPr.append(borders)
    p = cell.paragraphs[0]
    if lead:
        r = p.add_run(lead + " ")
        r.font.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(color)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string("23262B")
    return table


def add_totals_row(table, label, values, col_offset=2):
    """Bold shaded totals row: label in col 0, values from col_offset."""
    row = table.add_row().cells
    row[0].text = label
    for i, v in enumerate(values):
        if col_offset + i < len(row):
            row[col_offset + i].text = v
    for cell in row:
        shade_cell(cell, FMV_LIGHT_NAVY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(FMV_NAVY)
    for idx in range(col_offset, len(row)):
        for paragraph in row[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return row
