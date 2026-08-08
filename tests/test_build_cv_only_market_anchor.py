"""
Unit tests for scripts/build_cv_only_market_anchor.py.

Exercises:
  - interpolate_percentile() — piecewise-linear percentile mapping
  - blended_benchmark()      — n-weighted average across survey sources
  - build_exhibit()          — matplotlib PNG exhibit
  - build_report()           — DOCX 2-page report builder
  - main()                   — end-to-end exhibit + DOCX + JSON summary
"""
from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

import pytest

# Make scripts/ importable so the test can import the module directly.
REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "scripts"))

import build_cv_only_market_anchor as cvm  # noqa: E402


# ─── interpolate_percentile ────────────────────────────────────────────

class TestInterpolatePercentile:
    """Piecewise linear percentile mapping."""

    P25, P50, P75, P90 = 250000, 300000, 350000, 400000

    def test_below_p25(self) -> None:
        # At 0 → 0%; at p25 → 25%
        assert cvm.interpolate_percentile(0, self.P25, self.P50, self.P75, self.P90) == 0.0
        # At half of p25 → 12.5%
        result = cvm.interpolate_percentile(125000, self.P25, self.P50, self.P75, self.P90)
        assert result == pytest.approx(12.5)

    def test_at_p25(self) -> None:
        # At p25 exactly → 25%
        result = cvm.interpolate_percentile(self.P25, self.P25, self.P50, self.P75, self.P90)
        assert result == pytest.approx(25.0)

    def test_midpoint_between_p25_and_p50(self) -> None:
        # Halfway between p25 and p50 → 37.5%
        mid = (self.P25 + self.P50) / 2
        result = cvm.interpolate_percentile(mid, self.P25, self.P50, self.P75, self.P90)
        assert result == pytest.approx(37.5)

    def test_at_p50(self) -> None:
        result = cvm.interpolate_percentile(self.P50, self.P25, self.P50, self.P75, self.P90)
        assert result == pytest.approx(50.0)

    def test_midpoint_between_p50_and_p75(self) -> None:
        mid = (self.P50 + self.P75) / 2
        result = cvm.interpolate_percentile(mid, self.P25, self.P50, self.P75, self.P90)
        assert result == pytest.approx(62.5)

    def test_at_p75(self) -> None:
        result = cvm.interpolate_percentile(self.P75, self.P25, self.P50, self.P75, self.P90)
        assert result == pytest.approx(75.0)

    def test_midpoint_between_p75_and_p90(self) -> None:
        # Halfway between p75 and p90 → 82.5% (75 + 7.5)
        mid = (self.P75 + self.P90) / 2
        result = cvm.interpolate_percentile(mid, self.P25, self.P50, self.P75, self.P90)
        assert result == pytest.approx(82.5)

    def test_at_p90(self) -> None:
        result = cvm.interpolate_percentile(self.P90, self.P25, self.P50, self.P75, self.P90)
        assert result == pytest.approx(90.0)

    def test_above_p90_extrapolates(self) -> None:
        """Above p90 the function returns 90 + 10·(value-p90)/p90."""
        # value = 1.5 * p90 → 90 + 10·0.5 = 95
        result = cvm.interpolate_percentile(
            1.5 * self.P90, self.P25, self.P50, self.P75, self.P90
        )
        assert result == pytest.approx(95.0)

    def test_returns_none_when_input_missing(self) -> None:
        assert cvm.interpolate_percentile(300000, 0, self.P50, self.P75, self.P90) is None
        assert cvm.interpolate_percentile(300000, None, self.P50, self.P75, self.P90) is None
        assert cvm.interpolate_percentile(300000, self.P25, self.P50, self.P75, None) is None

    def test_returns_none_for_zero_value_inputs(self) -> None:
        # Zero percentiles count as "missing" and yield None.
        assert cvm.interpolate_percentile(300000, 0, 0, 0, 0) is None


# ─── blended_benchmark ────────────────────────────────────────────────

class TestBlendedBenchmark:
    """Weighted-average percentile blend across survey sources."""

    @pytest.fixture
    def benchmarks(self) -> dict:
        # SC n=89, Survey 3 n=142, Survey 2 n=34 (synthetic TCC numbers)
        return {
            "tcc": {
                "Survey 1": {"n": 89, "25": 285000, "50": 325000,
                                    "75": 370000, "90": 420000},
                "Survey 2":       {"n": 34, "25": 275000, "50": 315000,
                                    "75": 360000, "90": 405000},
                "Survey 3":            {"n": 142, "25": 280000, "50": 320000,
                                    "75": 365000, "90": 410000},
            },
            "wrvu": {},
            "base": {},
            "tcc_per_wrvu": {},
        }

    def test_blends_three_sources_with_weights(self, benchmarks) -> None:
        blended = cvm.blended_benchmark(benchmarks, "tcc")
        # Manually compute expected p50:
        # weighted = 325000*89 + 315000*34 + 320000*142
        # total_n = 89+34+142 = 265
        # = (28925000 + 10710000 + 45440000) / 265
        # = 85075000 / 265 = 321037.74...
        n_total = 89 + 34 + 142
        weighted_p50 = (325000 * 89 + 315000 * 34 + 320000 * 142) / n_total
        assert blended["50"] == pytest.approx(weighted_p50)

    def test_returns_none_for_section_without_sources(self, benchmarks) -> None:
        """Empty section → all percentiles None."""
        blended = cvm.blended_benchmark(benchmarks, "wrvu")
        for pct in ("25", "50", "75", "90"):
            assert blended[pct] is None

    def test_ignores_zero_n_rows(self) -> None:
        benchmarks = {
            "tcc": {
                "Survey 1": {"n": 0, "25": 100, "50": 200,
                                    "75": 300, "90": 400},
                "Survey 3":            {"n": 100, "25": 100, "50": 200,
                                    "75": 300, "90": 400},
            },
            "wrvu": {}, "base": {}, "tcc_per_wrvu": {},
        }
        blended = cvm.blended_benchmark(benchmarks, "tcc")
        # Only Survey 3 contributes (n=0 rows excluded)
        assert blended["50"] == 200

    def test_ignores_zero_value_rows(self) -> None:
        benchmarks = {
            "tcc": {
                "Survey 1": {"n": 50, "25": 0, "50": 0,
                                    "75": 0, "90": 0},
                "Survey 3":            {"n": 50, "25": 100, "50": 200,
                                    "75": 300, "90": 400},
            },
            "wrvu": {}, "base": {}, "tcc_per_wrvu": {},
        }
        blended = cvm.blended_benchmark(benchmarks, "tcc")
        # Zero-value rows excluded; only Survey 3 contributes
        assert blended["50"] == 200

    def test_returns_dict_keys_for_all_four_percentiles(self, benchmarks) -> None:
        blended = cvm.blended_benchmark(benchmarks, "tcc")
        assert set(blended.keys()) == {"25", "50", "75", "90"}


# ─── build_exhibit ─────────────────────────────────────────────────────

class TestBuildExhibit:
    def test_writes_valid_png(self, tmp_path: Path) -> None:
        blended = {"25": 280000, "50": 320000, "75": 365000, "90": 410000}
        out = tmp_path / "exhibit.png"
        cvm.build_exhibit(
            "Pediatric Critical Care", blended,
            recommended_low=280000, recommended_high=320000,
            recommended_mid=300000, output_path=str(out),
        )
        assert out.exists()
        assert out.stat().st_size > 5000
        with open(out, "rb") as f:
            assert f.read(8) == b"\x89PNG\r\n\x1a\n"


# ─── build_report ──────────────────────────────────────────────────────

class TestBuildReport:
    @pytest.fixture
    def exhibit(self, tmp_path: Path) -> str:
        blended = {"25": 280000, "50": 320000, "75": 365000, "90": 410000}
        out = tmp_path / "exhibit.png"
        cvm.build_exhibit(
            "Pediatric Critical Care", blended,
            recommended_low=280000, recommended_high=320000,
            recommended_mid=300000, output_path=str(out),
        )
        return str(out)

    def test_writes_valid_docx(self, exhibit, tmp_path: Path) -> None:
        out = tmp_path / "report.docx"
        blended = {"25": 280000, "50": 320000, "75": 365000, "90": 410000}
        cvm.build_report(
            "Provider 1", "Pediatric Critical Care", 8,
            blended, (280000, 300000, 320000), exhibit, str(out),
        )
        assert out.exists()
        with zipfile.ZipFile(out) as zf:
            assert "word/document.xml" in zf.namelist()

    def test_contains_caveat_and_recommended_range(
        self, exhibit, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "report.docx"
        blended = {"25": 280000, "50": 320000, "75": 365000, "90": 410000}
        cvm.build_report(
            "Provider 1", "Pediatric Critical Care", 8,
            blended, (280000, 300000, 320000), exhibit, str(out),
        )
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # The caveat now renders as a styled callout table, not a plain paragraph
        table_text = "\n".join(
            cell.text for tbl in doc.tables for row in tbl.rows for cell in row.cells
        )
        all_text = all_text + "\n" + table_text
        assert "IMPORTANT CAVEAT" in all_text
        assert "$280,000" in all_text
        assert "$300,000" in all_text
        assert "$320,000" in all_text


# ─── main() end-to-end ─────────────────────────────────────────────────

class TestMainEndToEnd:
    def test_main_writes_exhibit_docx_and_summary(
        self, mock_survey_xlsx: Path, tmp_path: Path
    ) -> None:
        out_dir = tmp_path / "cv_out"
        argv = [
            "build_cv_only_market_anchor.py",
            "--name", "Provider 1",
            "--specialty", "Pediatric Critical Care",
            "--yoe", "8",
            "--survey-file", str(mock_survey_xlsx),
            "--output-dir", str(out_dir),
            "--no-pdf",
        ]
        original_argv = sys.argv
        try:
            sys.argv = argv
            cvm.main()
        finally:
            sys.argv = original_argv

        assert (out_dir / "market_anchor_exhibit.png").exists()
        assert (out_dir / "market_anchor_report.docx").exists()
        assert (out_dir / "market_anchor_summary.json").exists()

        summary = json.loads(
            (out_dir / "market_anchor_summary.json").read_text()
        )
        assert summary["name"] == "Provider 1"
        assert summary["specialty"] == "Pediatric Critical Care"
        assert summary["yoe"] == 8
        # blended_tcc is dict with 4 keys
        assert set(summary["blended_tcc"].keys()) == {"25", "50", "75", "90"}
        # Recommended range has three values
        for k in ("low", "mid", "high"):
            assert k in summary["recommended_range"]
            assert summary["recommended_range"][k] > 0

    def test_main_with_benchmarks_text_file(
        self, tmp_path: Path
    ) -> None:
        """--benchmarks (text paste) path should also work end-to-end."""
        bench_file = tmp_path / "benchmarks.txt"
        bench_file.write_text(
            "TOTAL CASH COMPENSATION\n"
            "Survey 1\t89\t$285,000\t$325,000\t$370,000\t$420,000\n"
            "Survey 2\t34\t$275,000\t$315,000\t$360,000\t$405,000\n"
            "Survey 3\t142\t$280,000\t$320,000\t$365,000\t$410,000\n"
        )
        out_dir = tmp_path / "cv_out_text"
        argv = [
            "build_cv_only_market_anchor.py",
            "--name", "Provider 1",
            "--specialty", "Pediatric Critical Care",
            "--yoe", "8",
            "--benchmarks", str(bench_file),
            "--output-dir", str(out_dir),
            "--no-pdf",
        ]
        original_argv = sys.argv
        try:
            sys.argv = argv
            cvm.main()
        finally:
            sys.argv = original_argv
        assert (out_dir / "market_anchor_summary.json").exists()
