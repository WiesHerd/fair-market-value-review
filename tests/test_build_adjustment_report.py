"""
Unit tests for scripts/build_adjustment_report.py.

Exercises:
  - load_providers_tsv()   — TSV cohort parser
  - load_providers_excel() — Excel cohort parser
  - load_market_anchors()  — JSON anchors loader
  - build_exhibit()        — matplotlib PNG exhibit
  - build_report()         — DOCX 8-section report builder
  - main()                 — end-to-end CSV + DOCX + exhibit generation
"""
from __future__ import annotations

import csv
import json
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest

# Make scripts/ importable so the test can import the module directly.
REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "scripts"))

import build_adjustment_report as bar  # noqa: E402


# ─── TSV / Excel / JSON loaders ────────────────────────────────────────

class TestLoadProvidersTsv:
    def test_parses_three_providers(self, mock_cohort_tsv: Path) -> None:
        providers = bar.load_providers_tsv(str(mock_cohort_tsv))
        assert len(providers) == 3
        names = {p["name"] for p in providers}
        assert "Provider 1" in names
        assert "Provider 2" in names

    def test_parses_numeric_fields(self, mock_cohort_tsv: Path) -> None:
        providers = bar.load_providers_tsv(str(mock_cohort_tsv))
        jordan = next(p for p in providers if p["name"] == "Provider 1")
        assert jordan["yoe"] == 8
        assert jordan["current_base"] == 275000
        assert jordan["planned_base"] == 287000
        assert jordan["proposed_base"] == 325000

    def test_skips_short_rows(self, tmp_path: Path) -> None:
        f = tmp_path / "bad.tsv"
        f.write_text(
            "Name\tYOE\tCurrentBase\tPlannedBase\tProposedBase\n"
            "Only, Two\t5\t100\t200\n"  # only 4 columns
            "Real, Provider\t5\t100\t200\t300\n"
        )
        providers = bar.load_providers_tsv(str(f))
        assert len(providers) == 1
        assert providers[0]["name"] == "Real, Provider"

    def test_empty_yoe_stays_none(self, tmp_path: Path) -> None:
        f = tmp_path / "blank_yoe.tsv"
        f.write_text(
            "Name\tYOE\tCurrentBase\tPlannedBase\tProposedBase\n"
            "Doe, John\t\t100\t110\t120\n"
        )
        providers = bar.load_providers_tsv(str(f))
        assert providers[0]["yoe"] is None


class TestLoadProvidersExcel:
    def test_reads_mock_workbook(
        self, mock_salary_xlsx: Path, tmp_path: Path
    ) -> None:
        from tests.fixtures.builders import build_mock_salary_xlsx

        # Build a salary xlsx that matches the Excel-loader's column layout
        # (name_col=D, yoe_col=K, current_col=P, planned_col=Q, proposed_col=R).
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        # Header row (skipped by the loader starting at row 2)
        ws["D1"] = "Name"
        # Row 2 — provider A
        ws["D2"] = "Provider 1"
        ws["K2"] = 8
        ws["P2"] = 287000
        ws["Q2"] = 295000
        ws["R2"] = 325000
        xlsx = tmp_path / "salary_for_loader.xlsx"
        wb.save(xlsx)

        providers = bar.load_providers_excel(str(xlsx))
        assert len(providers) == 1
        assert providers[0]["name"] == "Provider 1"
        assert providers[0]["current_base"] == 287000

    def test_default_proposed_is_zero(self, tmp_path: Path) -> None:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws["D2"] = "Doe, Jane"
        ws["K2"] = 5
        ws["P2"] = 100
        ws["Q2"] = 110
        xlsx = tmp_path / "no_proposed.xlsx"
        wb.save(xlsx)
        providers = bar.load_providers_excel(str(xlsx))
        assert providers[0]["proposed_base"] == 0


class TestLoadMarketAnchors:
    def test_reads_p25_p50_p75_p90(self, mock_anchors_json: Path) -> None:
        anchors = bar.load_market_anchors(str(mock_anchors_json))
        assert anchors["p25"] == 280000
        assert anchors["p50"] == 320000
        assert anchors["p75"] == 365000
        assert anchors["p90"] == 410000


# ─── build_exhibit ─────────────────────────────────────────────────────

class TestBuildExhibit:
    def test_writes_png(self, tmp_path: Path) -> None:
        providers = [
            {"name": "Provider 1", "current_base": 287000,
             "planned_base": 295000, "proposed_base": 325000, "yoe": 8},
            {"name": "Provider 2", "current_base": 295000,
             "planned_base": 303000, "proposed_base": 330000, "yoe": 12},
        ]
        anchors = {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
        out = tmp_path / "exhibit.png"
        result = bar.build_exhibit(providers, anchors, str(out))
        assert result == str(out)
        assert out.exists()
        assert out.stat().st_size > 5000  # real PNG content

    def test_png_is_valid(self, tmp_path: Path) -> None:
        """The output should be a real PNG (magic number check)."""
        providers = [{
            "name": "Provider 1", "current_base": 287000,
            "planned_base": 295000, "proposed_base": 325000, "yoe": 8,
        }]
        anchors = {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
        out = tmp_path / "exhibit.png"
        bar.build_exhibit(providers, anchors, str(out))
        with open(out, "rb") as f:
            magic = f.read(8)
        assert magic[:8] == b"\x89PNG\r\n\x1a\n"

    def test_handles_missing_p90(self, tmp_path: Path) -> None:
        """p90 absent → no triangle marker, but chart still renders."""
        providers = [{
            "name": "Doe, Jane", "current_base": 100000,
            "planned_base": 105000, "proposed_base": 110000, "yoe": 3,
        }]
        anchors = {"p25": 80000, "p50": 100000, "p75": 120000}  # no p90
        out = tmp_path / "exhibit.png"
        bar.build_exhibit(providers, anchors, str(out))
        assert out.exists()


# ─── build_report ──────────────────────────────────────────────────────

class TestBuildReport:
    @pytest.fixture
    def exhibit(self, tmp_path: Path) -> str:
        providers = [{
            "name": "Provider 1", "current_base": 287000,
            "planned_base": 295000, "proposed_base": 325000, "yoe": 8,
        }]
        anchors = {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
        out = tmp_path / "exhibit.png"
        bar.build_exhibit(providers, anchors, str(out))
        return str(out)

    @pytest.fixture
    def providers(self) -> list:
        return [
            {"name": "Provider 1", "yoe": 8,
             "current_base": 287000, "planned_base": 295000,
             "proposed_base": 325000},
            {"name": "Provider 2", "yoe": 12,
             "current_base": 295000, "planned_base": 303000,
             "proposed_base": 330000},
        ]

    def test_writes_valid_docx(
        self, providers, exhibit, tmp_path: Path
    ) -> None:
        out = tmp_path / "report.docx"
        anchors = {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
        bar.build_report(providers, anchors, exhibit, str(out))
        assert out.exists()
        # DOCX is a ZIP with word/document.xml
        with zipfile.ZipFile(out) as zf:
            names = zf.namelist()
        assert "word/document.xml" in names

    def test_contains_eight_section_headings(
        self, providers, exhibit, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "report.docx"
        anchors = {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
        bar.build_report(providers, anchors, exhibit, str(out))
        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
        expected = [
            "Fair Market Value Review",  # title (level 0)
            "Data Sources",
            "Executive Summary",
            "Cohort Detail",
            "Market Range Exhibit",
            "Logic & Empirical Rationale",
            "Cost Impact",
            "Governance Notes & Guardrails",
            "Footnotes & Caveats",
        ]
        for h in expected:
            assert h in headings, f"Missing heading: {h}"

    def test_includes_cohort_size_in_summary(
        self, providers, exhibit, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "report.docx"
        anchors = {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
        bar.build_report(providers, anchors, exhibit, str(out))
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Cohort size: 2 provider" in all_text

    def test_includes_proposed_total(
        self, providers, exhibit, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "report.docx"
        anchors = {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
        bar.build_report(providers, anchors, exhibit, str(out))
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # 325000 + 330000 = 655000
        assert "$655,000" in all_text

    def test_metadata_overrides_default_text(
        self, providers, exhibit, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "report.docx"
        anchors = {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
        meta = {
            "cohort_name": "Pediatric Critical Care Faculty",
            "division": "Critical Care",
            "rationale": ["Custom rationale line 1", "Custom rationale line 2"],
        }
        bar.build_report(providers, anchors, exhibit, str(out), metadata=meta)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Pediatric Critical Care Faculty" in all_text
        assert "Custom rationale line 1" in all_text

    def test_benefits_rate_adds_loaded_cost_paragraph(
        self, providers, exhibit, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "report.docx"
        anchors = {"p25": 280000, "p50": 320000, "p75": 365000, "p90": 410000}
        meta = {"benefits_rate": 0.20}
        bar.build_report(providers, anchors, exhibit, str(out), metadata=meta)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "20.0%" in all_text


# ─── main() end-to-end ─────────────────────────────────────────────────

class TestMainEndToEnd:
    def test_main_generates_docx_csv_and_exhibit(
        self, mock_cohort_tsv, mock_anchors_json, tmp_path: Path
    ) -> None:
        out_docx = tmp_path / "report.docx"
        argv = [
            "build_adjustment_report.py",
            "--name", "Pediatric Critical Care Faculty",
            "--providers-file", str(mock_cohort_tsv),
            "--market-anchors-file", str(mock_anchors_json),
            "--output", str(out_docx),
            "--no-pdf",  # skip LibreOffice headless in CI
        ]
        original_argv = sys.argv
        try:
            sys.argv = argv
            bar.main()
        finally:
            sys.argv = original_argv

        assert out_docx.exists()
        # CSV companion
        csv_path = out_docx.with_suffix(".csv")
        assert csv_path.exists()
        with open(csv_path) as f:
            rows = list(csv.DictReader(f))
        assert len(rows) == 3
        assert rows[0]["Name"] == "Provider 1"
        # Exhibit
        exhibit_path = tmp_path / "report_exhibit.png"
        assert exhibit_path.exists()

    def test_main_csv_export_columns(
        self, mock_cohort_tsv, mock_anchors_json, tmp_path: Path
    ) -> None:
        out_docx = tmp_path / "report.docx"
        argv = [
            "build_adjustment_report.py",
            "--providers-file", str(mock_cohort_tsv),
            "--market-anchors-file", str(mock_anchors_json),
            "--output", str(out_docx),
            "--no-pdf",
        ]
        original_argv = sys.argv
        try:
            sys.argv = argv
            bar.main()
        finally:
            sys.argv = original_argv
        with open(out_docx.with_suffix(".csv")) as f:
            reader = csv.reader(f)
            header = next(reader)
        assert header == [
            "Name", "YOE", "CurrentBase", "PlannedBase",
            "ProposedBase", "DeltaVsCurrent", "DeltaVsPlanned",
        ]

    def test_main_missing_providers_file_exits(self) -> None:
        argv = ["build_adjustment_report.py"]
        original_argv = sys.argv
        try:
            sys.argv = argv
            with pytest.raises(SystemExit) as exc:
                bar.main()
        finally:
            sys.argv = original_argv
        assert exc.value.code == 1
